"""
Extracts raw text from an RFA/NOFO, regardless of source: uploaded PDF,
uploaded DOCX, a pasted URL, or pasted text.
"""
from __future__ import annotations

import io
import re

import pdfplumber
from docx import Document

from scrape_grant_page import scrape


class ExtractionError(Exception):
    """Raised when text could not be extracted from the given source."""


def extract_text_from_pdf(file_bytes: bytes) -> str:
    pages = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            for table in page.extract_tables() or []:
                for row in table:
                    cells = [c.strip() for c in row if c and c.strip()]
                    if cells:
                        page_text += "\n• " + " | ".join(cells)
            if page_text:
                lines = page_text.splitlines()
                page_text = "\n".join(
                    l for l in lines
                    if not re.match(r"^\s*(page\s*\d+|\d+\s*of\s*\d+)\s*$", l, re.I)
                )
            pages.append(page_text.strip())

    text = "\n\n".join(p for p in pages if p).strip()
    if not text:
        raise ExtractionError(
            "PDF appears to be scanned or image-only — no text could be extracted."
        )
    return text


def extract_text_from_docx(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append("• " + " | ".join(cells))

    text = "\n".join(parts).strip()
    if not text:
        raise ExtractionError("DOCX contained no extractable text.")
    return text


def extract_text_from_url(url: str) -> str:
    text = scrape(url)
    if not text.strip():
        raise ExtractionError("No text could be extracted from that URL.")
    return text


def extract_text(*, pdf_bytes: bytes | None = None, docx_bytes: bytes | None = None,
                  url: str | None = None, pasted_text: str | None = None) -> str:
    """Dispatch to the right extractor based on whichever input was provided."""
    if pasted_text and pasted_text.strip():
        return pasted_text.strip()
    if pdf_bytes:
        return extract_text_from_pdf(pdf_bytes)
    if docx_bytes:
        return extract_text_from_docx(docx_bytes)
    if url:
        return extract_text_from_url(url)
    raise ExtractionError("No RFA source provided (upload a file, paste a URL, or paste text).")
