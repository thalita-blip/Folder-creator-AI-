"""
Renders a synthesized SOP JSON object (see tools/synthesize.py) into a
formatted SOP.docx matching the fixed 16-section structure.
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt

SECTION_TITLES = [
    ("grant_purpose", "Grant Purpose"),
    ("grant_overview", "Grant Overview"),
    ("reimbursement_rules", "Reimbursement Rules"),
    ("eligible_projects", "Eligible Projects"),
    ("eligible_costs", "Eligible Costs"),
    ("ineligible_costs", "Ineligible Costs"),
    ("most_important_rules", "Most Important Grant Rules"),
    ("applicant_eligibility_questions", "Applicant Eligibility Questions"),
    ("project_eligibility_questions", "Project Eligibility Questions"),
    ("evaluation_criteria", "Evaluation Criteria"),
    ("strengthening_questions", "Questions That Strengthen the Application"),
    ("project_readiness_checklist", "Project Readiness Checklist"),
    ("documents_to_request", "Documents to Request"),
    ("financing_questions", "Financing Questions"),
    ("client_commitment_questions", "Client Commitment Questions"),
    ("red_flags", "Red Flags"),
]

_GROUP_LABELS = {
    "business": "Business",
    "previous_funding": "Previous Funding",
    "compliance": "Compliance",
    "financial": "Financial",
    "project": "Project",
    "supporting": "Supporting",
}


def _add_bullets(doc: Document, items: list[str]):
    for item in items:
        doc.add_paragraph(str(item), style="List Bullet")


def _add_grouped_bullets(doc: Document, grouped: dict):
    for key, items in grouped.items():
        if not items:
            continue
        label = _GROUP_LABELS.get(key, key.replace("_", " ").title())
        doc.add_paragraph(label, style="Heading 3")
        _add_bullets(doc, items)


def _add_evaluation_criteria(doc: Document, criteria: list[dict]):
    for entry in criteria:
        points = entry.get("points")
        header = entry.get("criterion", "")
        if points is not None:
            header = f"{header} — {points} pts"
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(header)
        run.bold = True
        if entry.get("notes"):
            doc.add_paragraph(entry["notes"])


def _render_section_body(doc: Document, key: str, value):
    if value is None or value == "" or value == []:
        doc.add_paragraph("Not addressed in this RFA.")
        return

    if key == "evaluation_criteria":
        _add_evaluation_criteria(doc, value)
    elif key in ("applicant_eligibility_questions", "documents_to_request"):
        _add_grouped_bullets(doc, value)
    elif isinstance(value, list):
        _add_bullets(doc, value)
    else:
        doc.add_paragraph(str(value))


def generate_sop_docx(grant_name: str, year: str, sop_sections: dict, output_path: Path) -> Path:
    doc = Document()

    title = doc.add_heading(f"{grant_name} ({year}) — SOP", level=0)

    for key, title_text in SECTION_TITLES:
        doc.add_heading(title_text, level=1)
        _render_section_body(doc, key, sop_sections.get(key))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
